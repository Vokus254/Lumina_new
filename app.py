import io
import html
import uuid
import re
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st

try:
    from supabase import create_client
except Exception:
    create_client = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from mapping import get_dynamic_mapping
except Exception:
    get_dynamic_mapping = None


# ============================================================
# LUMINA Mapping – internes Abschluss-Cockpit
# Streamlit + optional Supabase
# ============================================================

st.set_page_config(
    page_title="LUMINA Mapping – Abschluss-Cockpit",
    page_icon="📊",
    layout="wide",
)

KLARUNG = "9. KLÄRUNGSPOSTEN"
APP_VERSION = "2026-05-10"
APP_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = APP_DIR / "templates"
DEFAULT_MAPPING_NAME = "Standard"
OPENAI_MODELS = {
    "GPT-5.2 (Standard)": "gpt-5.2",
    "GPT-5": "gpt-5",
    "GPT-4.1": "gpt-4.1",
    "GPT-4.1 mini": "gpt-4.1-mini",
    "GPT-4.1 nano": "gpt-4.1-nano",
    "GPT-4o": "gpt-4o",
    "GPT-4o mini": "gpt-4o-mini",
    "GPT-4 (älter)": "gpt-4",
}
AI_ROLES = [
    "Wirtschaftsprüfer",
    "HGB-Bilanzierer",
    "Management-Reporter",
    "CFO",
    "Leiter Rechnungswesen",
]
AI_PURPOSES = [
    "HGB-Anhang",
    "Lagebericht",
    "Management Summary",
    "Prüfer-Rückfragen",
    "Bank-Reporting",
    "Internes Rechnungswesen",
]
AI_AUDIENCES = [
    "Geschäftsführung",
    "Wirtschaftsprüfer",
    "Bank",
    "Internes Rechnungswesen",
    "Beirat / Aufsichtsorgan",
]
AI_TONES = [
    "prüferfreundlich, vorsichtig, sachlich, konservativ",
    "managementorientiert, klar, entscheidungsnah",
    "bankorientiert, risikobewusst, faktenbasiert",
    "intern, knapp, handlungsorientiert",
]
AI_LENGTHS = ["kurz", "mittel", "lang"]
AI_FORMS = ["gegliederte Abschnitte", "Bulletpoints", "Fließtext", "Tabelle"]
MANDANT_FIELDS = [
    "mandantenname",
    "rechtsform",
    "branche",
    "sitz",
    "geschaeftsjahr",
    "kontenrahmen",
    "groessenklasse",
    "pruefungspflicht",
    "lageberichtspflicht",
    "steuerberater",
    "wirtschaftspruefer",
    "ansprechpartner_rechnungswesen",
    "besonderheiten",
]
ONBOARDING_SECTIONS = {
    "A. Stammdaten": ["Welche Gesellschaften/Einheiten sind im Abschluss enthalten?", "Gibt es Änderungen im Konsolidierungs- oder Berichtskreis?"],
    "B. Rechtsform und Größenklasse": ["Welche Größenklasse liegt vor?", "Bestehen besondere Offenlegungs- oder Prüfungspflichten?"],
    "C. Branche und Geschäftsmodell": ["Welche wesentlichen Erlösquellen bestehen?", "Welche operativen Besonderheiten prägen das Geschäftsjahr?"],
    "D. Abschlussbesonderheiten": ["Gab es Umstellungen, Sondereffekte oder bilanzpolitische Entscheidungen?", "Welche Sachverhalte sollen im Reporting besonders erklärt werden?"],
    "E. Forderungen / OP / Wertberichtigungen": ["Gibt es überfällige oder risikobehaftete Forderungen?", "Wie wurden Wertberichtigungen ermittelt?"],
    "F. Anlagevermögen": ["Gab es wesentliche Zugänge, Abgänge oder außerplanmäßige Abschreibungen?", "Sind Investitionsprojekte oder Anlagen im Bau wesentlich?"],
    "G. Rückstellungen": ["Welche wesentlichen Rückstellungsarten bestehen?", "Gab es Auflösungen, Inanspruchnahmen oder Neubildungen?"],
    "H. Verbindlichkeiten / Darlehen": ["Welche Darlehen oder Finanzierungen sind wesentlich?", "Gab es neue Covenants, Tilgungen oder Umschuldungen?"],
    "I. Intercompany / Verrechnungskonten": ["Welche konzerninternen Salden bestehen?", "Sind Verrechnungskonten abgestimmt und dokumentiert?"],
    "J. Reporting-Stil": ["Welche Tonalität erwartet die Zielgruppe?", "Welche Kennzahlen oder Schwerpunkte sollen priorisiert werden?"],
    "K. Dokumentenanforderungen": ["Welche Nachweise werden für Abschluss und Prüfung benötigt?", "Welche Dokumente fehlen aktuell noch?"],
}
SUSA_TYPES = ["Hauptgesellschaft", "Verwaltung", "Betriebsstätte", "Einrichtung", "Teilbereich", "Konsolidierte SuSa"]


# ------------------------------------------------------------
# Styling
# ------------------------------------------------------------
st.markdown(
    """
    <style>
    .main .block-container {padding-top: 1.5rem; padding-bottom: 2rem;}
    .lumina-box {
        border: 1px solid #e6e6e6;
        border-radius: 14px;
        padding: 1rem 1.2rem;
        background: #fafafa;
        margin-bottom: 1rem;
    }
    .small-muted {color: #666; font-size: 0.9rem;}
    </style>
    """,
    unsafe_allow_html=True,
)


# ------------------------------------------------------------
# Helpers
# ------------------------------------------------------------
def normalize_konto(value) -> str:
    """Normalisiert Kontonummern robust: 8400.0 -> 8400, Leerzeichen raus."""
    if pd.isna(value):
        return ""
    s = str(value).strip()
    s = re.sub(r"\s+", "", s)
    if s.endswith(".0"):
        s = s[:-2]
    s = re.sub(r"[^0-9A-Za-z_-]", "", s)
    return s


def parse_number(value) -> float:
    """Konvertiert deutsche/englische Zahlenformate in float."""
    if pd.isna(value):
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)

    s = str(value).strip().replace("\u00a0", "")
    if s in ["", "-", "None", "nan"]:
        return 0.0

    negative = False
    if s.startswith("(") and s.endswith(")"):
        negative = True
        s = s[1:-1]
    if s.endswith("-"):
        negative = True
        s = s[:-1]

    s = (
        s.replace("€", "")
        .replace("EUR", "")
        .replace(" ", "")
        .replace("'", "")
    )

    # deutsches Format: 1.234,56
    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")

    try:
        number = float(s)
        return -number if negative else number
    except Exception:
        return 0.0


def format_de_number(value) -> str:
    """Formatiert Zahlen für die App-Anzeige im deutschen Format: 10.005,56."""
    if pd.isna(value):
        return ""
    try:
        return f"{float(value):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return str(value)


def format_de_amount(value, suffix: str = "") -> str:
    formatted = format_de_number(value)
    return f"{formatted} {suffix}".strip()


def display_dataframe(df: pd.DataFrame, value_cols: list[str] | None = None):
    if value_cols is None:
        value_cols = detect_value_cols(df)
    formatters = {c: format_de_number for c in value_cols if c in df.columns}
    if not formatters:
        return df
    return df.style.format(formatters)


def normalize_column_name(value) -> str:
    if isinstance(value, (datetime, pd.Timestamp)):
        return value.strftime("%d.%m.%Y")
    text = str(value).strip()
    parsed = pd.to_datetime(text, errors="coerce")
    if not pd.isna(parsed) and re.fullmatch(r"\d{4}-\d{2}-\d{2}( 00:00:00)?", text):
        return parsed.strftime("%d.%m.%Y")
    return text


def read_excel_smart(uploaded_file) -> pd.DataFrame:
    """Liest Excel und sucht die wahrscheinlich richtige Kopfzeile."""
    raw = pd.read_excel(uploaded_file, header=None, dtype=object)
    header_row = 0

    for i, row in raw.head(30).iterrows():
        row_text = " | ".join(row.dropna().astype(str).tolist()).lower()
        if any(x in row_text for x in ["konto", "kontonummer", "konto-nr", "account"]):
            header_row = i
            break

    uploaded_file.seek(0)
    df = pd.read_excel(uploaded_file, header=header_row, dtype=object)
    df = df.dropna(how="all")
    df.columns = [normalize_column_name(c) for c in df.columns]
    return df


def first_col(df: pd.DataFrame, contains: list[str]) -> str | None:
    for c in df.columns:
        name = str(c).lower()
        if all(x.lower() in name for x in contains):
            return c
    return None


def detect_konto_col(df: pd.DataFrame) -> str | None:
    candidates = []
    for c in df.columns:
        name = str(c).lower()
        if "konto" in name or "account" in name or name in {"kt", "kto", "kontonr"}:
            candidates.append(c)
    return candidates[0] if candidates else None


def detect_text_col(df: pd.DataFrame) -> str | None:
    for key in ["kontobezeichnung", "bezeichnung", "konto text", "text", "name"]:
        c = first_col(df, [key])
        if c:
            return c
    return None


def detect_value_cols(df: pd.DataFrame) -> list[str]:
    value_cols = []
    for c in df.columns:
        name = str(c).lower()
        if any(x in name for x in ["saldo", "betrag", "wert", "summe", "202", "31.12", "haben", "soll", "debit", "credit", "balance"]):
            # Spalten mit Konto-/Bezeichnung nicht versehentlich als Wert nehmen
            if "konto" not in name and "bezeichnung" not in name and "text" not in name:
                value_cols.append(c)

    # Fallback: numerische Spalten
    if not value_cols:
        for c in df.columns:
            converted = df[c].apply(parse_number)
            if converted.abs().sum() != 0:
                value_cols.append(c)

    return value_cols


def normalize_susa(df: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
    konto_col = detect_konto_col(df)
    text_col = detect_text_col(df)
    value_cols = detect_value_cols(df)

    if not konto_col:
        raise ValueError("Keine Kontospalte gefunden. Bitte Spalte z. B. 'Konto' oder 'Kontonummer' nennen.")
    if not value_cols:
        raise ValueError("Keine Wert-/Saldo-Spalte gefunden. Bitte Spalte z. B. 'Saldo 2025' nennen.")

    out = df.copy()
    out["KontoNr"] = out[konto_col].apply(normalize_konto)
    out = out[out["KontoNr"] != ""].copy()

    if text_col:
        out["Kontobezeichnung"] = out[text_col].astype(str).fillna("")
    else:
        out["Kontobezeichnung"] = ""

    for c in value_cols:
        out[c] = out[c].apply(parse_number)

    meta = {
        "konto_col": konto_col,
        "text_col": text_col,
        "value_cols": value_cols,
    }
    return out, meta


def is_unassigned_mapping(row: pd.Series) -> bool:
    values = [str(row.get(f"Ausweis_{i}", "")).strip() for i in range(1, 8)]
    return all(v == "" or v == KLARUNG or v.lower() == "nicht zugeordnet" for v in values)


def is_clarification_value(value) -> bool:
    return str(value).strip() == KLARUNG


def clarification_mask(df: pd.DataFrame) -> pd.Series:
    mask = df.get("Mapping_Status", pd.Series("", index=df.index)).eq("Klärung")
    if "Ausweis_1" in df.columns:
        mask = mask | df["Ausweis_1"].apply(is_clarification_value)
    if any(f"Ausweis_{i}" in df.columns for i in range(1, 8)):
        mask = mask | df.apply(is_unassigned_mapping, axis=1)
    return mask


def mapping_counts(df: pd.DataFrame) -> dict:
    klarung = clarification_mask(df)
    return {
        "total": len(df),
        "klarung": int(klarung.sum()),
        "vorschlag": int((df["Mapping_Status"] == "Vorschlag").sum()) if "Mapping_Status" in df.columns else 0,
        "gemappt": int(((df["Mapping_Status"] == "gemappt") & ~klarung).sum()) if "Mapping_Status" in df.columns else int((~klarung).sum()),
    }


def builtin_mapping_for(konto_nr: str) -> dict:
    """Liefert die lokale HGB-Regel aus mapping.py, falls eine echte Zuordnung existiert."""
    if get_dynamic_mapping is None:
        return {}
    try:
        candidate = get_dynamic_mapping(konto_nr)
    except Exception:
        return {}
    if not candidate:
        return {}

    normalized = {f"Ausweis_{i}": str(candidate.get(f"Ausweis_{i}", "") or "").strip() for i in range(1, 8)}
    if all(v == "" or v.lower() == "nicht zugeordnet" for v in normalized.values()):
        return {}
    return normalized


def empty_mapping_frame() -> pd.DataFrame:
    return pd.DataFrame(columns=["KontoNr"] + [f"Ausweis_{i}" for i in range(1, 8)])


def clean_manual_mapping_value(value) -> str:
    text = "" if pd.isna(value) else str(value).strip()
    return "" if text == KLARUNG else text


def rows_to_mapping_updates(rows: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    updates = []
    skipped = []
    for _, row in rows.iterrows():
        konto = normalize_konto(row.get("KontoNr"))
        if not konto:
            continue
        item = {"KontoNr": konto}
        for i in range(1, 8):
            item[f"Ausweis_{i}"] = clean_manual_mapping_value(row.get(f"Ausweis_{i}", ""))
        if not item["Ausweis_1"]:
            skipped.append(konto)
            continue
        updates.append(item)
    return pd.DataFrame(updates, columns=["KontoNr"] + [f"Ausweis_{i}" for i in range(1, 8)]), skipped


def merge_mapping_updates(current_mapping: pd.DataFrame | None, updates: pd.DataFrame) -> pd.DataFrame:
    base = normalize_mapping(current_mapping) if current_mapping is not None else empty_mapping_frame()
    if updates.empty:
        return base
    return normalize_mapping(pd.concat([base, updates], ignore_index=True))


def ausweis_options(mapping: pd.DataFrame | None, level: int, extra_values: pd.Series | None = None) -> list[str]:
    col = f"Ausweis_{level}"
    options = {""}
    if mapping is not None and col in mapping.columns:
        for value in mapping[col].dropna().astype(str).str.strip():
            if value and value != KLARUNG:
                options.add(value)
    if extra_values is not None:
        for value in extra_values.dropna().astype(str).str.strip():
            if value and value != KLARUNG:
                options.add(value)
    return [""] + sorted(options - {""})


def ausweis_column_config(mapping: pd.DataFrame | None, editor_df: pd.DataFrame | None = None) -> dict:
    return {
        f"Ausweis_{i}": st.column_config.SelectboxColumn(
            f"Ausweis_{i}",
            options=ausweis_options(
                mapping,
                i,
                editor_df[f"Ausweis_{i}"] if editor_df is not None and f"Ausweis_{i}" in editor_df.columns else None,
            ),
            required=(i == 1),
        )
        for i in range(1, 8)
    }


def normalize_mapping(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out.columns = [str(c).strip() for c in out.columns]

    # Konto-Spalte harmonisieren
    konto_col = None
    for c in out.columns:
        name = str(c).lower()
        if name in ["kontonr", "konto", "konto_nr", "kontonummer", "account"] or "konto" in name:
            konto_col = c
            break
    if not konto_col:
        raise ValueError("Im Mapping fehlt eine Konto-Spalte, z. B. 'KontoNr'.")

    out["KontoNr"] = out[konto_col].apply(normalize_konto)

    # Ausweis-Spalten harmonisieren
    rename = {}
    for c in out.columns:
        name = str(c).lower().replace(" ", "_")
        m = re.search(r"ausweis[_-]?(\d)", name)
        if m:
            rename[c] = f"Ausweis_{m.group(1)}"
    out = out.rename(columns=rename)

    for i in range(1, 8):
        col = f"Ausweis_{i}"
        if col not in out.columns:
            out[col] = ""
        out[col] = out[col].fillna("").astype(str).str.strip()

    assigned_mask = ~out["Ausweis_1"].isin(["", KLARUNG])
    for i in range(2, 8):
        col = f"Ausweis_{i}"
        out.loc[assigned_mask & (out[col] == KLARUNG), col] = ""

    out = out[out["KontoNr"] != ""].copy()
    out = out.drop_duplicates(subset=["KontoNr"], keep="last")
    return out[["KontoNr"] + [f"Ausweis_{i}" for i in range(1, 8)]]


def apply_mapping(susa: pd.DataFrame, mapping: pd.DataFrame) -> pd.DataFrame:
    map_df = normalize_mapping(mapping)
    df = susa.copy()
    df["KontoNr"] = df["KontoNr"].apply(normalize_konto)

    mapped = df.merge(map_df, on="KontoNr", how="left", indicator=True)
    mapped["Mapping_Status"] = mapped["_merge"].map({"both": "gemappt", "left_only": "Klärung"}).astype(str)
    mapped = mapped.drop(columns=["_merge"])

    for i in range(1, 8):
        col = f"Ausweis_{i}"
        mapped[col] = mapped[col].fillna("").astype(str).str.strip()

    for idx, row in mapped.iterrows():
        needs_clarification = row["Mapping_Status"] == "Klärung" or is_unassigned_mapping(row)
        if not needs_clarification:
            for i in range(2, 8):
                col = f"Ausweis_{i}"
                if mapped.at[idx, col] == KLARUNG:
                    mapped.at[idx, col] = ""
            continue

        fallback = builtin_mapping_for(row["KontoNr"])
        if fallback:
            for i in range(1, 8):
                col = f"Ausweis_{i}"
                value = fallback.get(col, "")
                if value:
                    mapped.at[idx, col] = value
                elif mapped.at[idx, col] in ["", KLARUNG]:
                    mapped.at[idx, col] = "Vorschlag offen"
            mapped.at[idx, "Mapping_Status"] = "Vorschlag"
            continue

        for i in range(1, 8):
            col = f"Ausweis_{i}"
            if mapped.at[idx, col] == "":
                mapped.at[idx, col] = KLARUNG
        mapped.at[idx, "Mapping_Status"] = "Klärung"

    return mapped


def build_pivot(df: pd.DataFrame, group_level: int, value_cols: list[str]) -> pd.DataFrame:
    ausweis_cols = [f"Ausweis_{i}" for i in range(1, group_level + 1) if f"Ausweis_{i}" in df.columns]
    if not ausweis_cols:
        return pd.DataFrame()
    return df.groupby(ausweis_cols, dropna=False)[value_cols].sum().reset_index()


def analysis_columns(value_cols: list[str]) -> tuple[str | None, str | None]:
    if not value_cols:
        return None, None
    current_col = value_cols[0]
    prior_col = value_cols[1] if len(value_cols) > 1 else None
    return current_col, prior_col


def add_variance_columns(df: pd.DataFrame, current_col: str, prior_col: str | None) -> pd.DataFrame:
    out = df.copy()
    out["Laufendes Jahr"] = out[current_col]
    out["Vorjahr"] = out[prior_col] if prior_col else 0.0
    out["Veränderung"] = out["Laufendes Jahr"] - out["Vorjahr"]
    out["Veränderung_abs"] = out["Veränderung"].abs()
    out["Veränderung_%"] = out.apply(
        lambda r: (r["Veränderung"] / abs(r["Vorjahr"]) * 100) if r["Vorjahr"] else pd.NA,
        axis=1,
    )
    return out


def grouped_variance(df: pd.DataFrame, group_cols: list[str], current_col: str, prior_col: str | None) -> pd.DataFrame:
    usable_group_cols = [c for c in group_cols if c in df.columns]
    if not usable_group_cols:
        return pd.DataFrame()
    work = df.copy()
    if prior_col is None:
        work["__prior"] = 0.0
        prior_col = "__prior"

    grouped = work.groupby(usable_group_cols, dropna=False)[[current_col, prior_col]].sum().reset_index()
    grouped = grouped.rename(columns={current_col: "Laufendes Jahr", prior_col: "Vorjahr"})
    grouped["Veränderung"] = grouped["Laufendes Jahr"] - grouped["Vorjahr"]
    grouped["Veränderung_abs"] = grouped["Veränderung"].abs()
    grouped["Veränderung_%"] = grouped.apply(
        lambda r: (r["Veränderung"] / abs(r["Vorjahr"]) * 100) if r["Vorjahr"] else pd.NA,
        axis=1,
    )
    return grouped.sort_values("Veränderung_abs", ascending=False)


def format_percent(value) -> str:
    if pd.isna(value):
        return "n/a"
    return f"{float(value):,.1f} %".replace(",", "X").replace(".", ",").replace("X", ".")


def clean_ai_output(text: str) -> str:
    cleaned = text or ""
    cleaned = cleaned.replace("**", "")
    cleaned = re.sub(r"^#{1,6}\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^\s*[-*]\s+", "- ", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def ai_output_to_docx(text: str, title: str) -> bytes | None:
    if Document is None:
        return None

    document = Document()
    document.add_heading(title, level=1)

    for block in clean_ai_output(text).splitlines():
        line = block.strip()
        if not line:
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif len(line) < 90 and not line.endswith("."):
            document.add_heading(line, level=2)
        else:
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def variance_display(df: pd.DataFrame):
    display_cols = [c for c in df.columns if c != "Veränderung_abs"]
    formatters = {
        "Laufendes Jahr": format_de_number,
        "Vorjahr": format_de_number,
        "Veränderung": format_de_number,
        "Veränderung_%": format_percent,
    }
    return df[display_cols].style.format({k: v for k, v in formatters.items() if k in df.columns})


def interpretation_markdown(
    df: pd.DataFrame,
    value_cols: list[str],
    mandant: str,
    abschlussjahr: int,
    threshold: float,
    top_n: int,
    ai_config: dict | None = None,
    mandant_profile: dict | None = None,
    reporting_profile: dict | None = None,
    onboarding_answers: dict | None = None,
) -> str:
    current_col, prior_col = analysis_columns(value_cols)
    if current_col is None:
        return "Keine Wertspalten gefunden."

    work = add_variance_columns(df, current_col, prior_col)
    level_summary = grouped_variance(df, ["Ausweis_1", "Ausweis_2", "Ausweis_3"], current_col, prior_col).head(top_n)
    account_cols = ["KontoNr", "Kontobezeichnung", "Ausweis_1", "Ausweis_2", "Ausweis_3"]
    account_cols = [c for c in account_cols if c in work.columns]
    top_accounts = work[work["Veränderung_abs"] >= threshold].sort_values("Veränderung_abs", ascending=False).head(top_n)

    ai_config = ai_config or {}
    mandant_profile = mandant_profile or {}
    reporting_profile = reporting_profile or {}
    onboarding_answers = onboarding_answers or {}
    rules = ai_config.get("rules", [])
    output_structure = ai_config.get("output_structure", [])

    lines = [
        f"# KI-Arbeitsgrundlage für {mandant} {abschlussjahr}",
        "",
        "## Steuerung der KI-Ausgabe",
        f"- Rolle: {ai_config.get('role', 'erfahrener HGB-Abschlussanalyst')}",
        f"- Zweck: {ai_config.get('purpose', 'Management-Reporting, Anhang-Hinweise und Lagebericht-Hinweise')}",
        f"- Adressat: {ai_config.get('audience', 'Geschäftsführung und Wirtschaftsprüfer')}",
        f"- Ton / Stil: {ai_config.get('tone', 'prüferfreundlich, vorsichtig, sachlich, konservativ')}",
        f"- Textumfang: {ai_config.get('length', 'mittel')}",
        f"- Form: {ai_config.get('form', 'gegliederte Abschnitte')}",
        "",
        "## Mandantenprofil",
        f"- Rechtsform: {mandant_profile.get('rechtsform', '')}",
        f"- Branche: {mandant_profile.get('branche', '')}",
        f"- Sitz: {mandant_profile.get('sitz', '')}",
        f"- Geschäftsjahr: {mandant_profile.get('geschaeftsjahr', '')}",
        f"- Kontenrahmen: {mandant_profile.get('kontenrahmen', '')}",
        f"- Größenklasse: {mandant_profile.get('groessenklasse', '')}",
        f"- Prüfungspflicht: {mandant_profile.get('pruefungspflicht', '')}",
        f"- Lageberichtspflicht: {mandant_profile.get('lageberichtspflicht', '')}",
        f"- Besonderheiten: {mandant_profile.get('besonderheiten', '')}",
        "",
        "## Reporting-Profil",
        f"- Berichtsstil: {reporting_profile.get('berichtsstil', '')}",
        f"- Zielgruppe: {reporting_profile.get('zielgruppe', '')}",
        f"- Anhang-Level: {reporting_profile.get('anhang_level', '')}",
        f"- Lagebericht-Stil: {reporting_profile.get('lagebericht_stil', '')}",
        "",
        "## Onboarding-Antworten",
        *(
            f"- {section} / {question}: {answer}"
            for (section, question), answer in onboarding_answers.items()
            if str(answer).strip()
        ),
        "",
        "## Kontext",
        f"- Abschlussjahr: {abschlussjahr}",
        f"- Laufendes Jahr: {current_col}",
        f"- Vorjahr: {prior_col or 'nicht vorhanden'}",
        f"- Wesentlichkeitsschwelle für Auffälligkeiten: {format_de_amount(threshold, 'EUR')}",
        "",
        "## Auffällige Abschlusspositionen",
    ]

    if level_summary.empty:
        lines.append("- Keine aggregierten Positionen auswertbar.")
    else:
        for _, row in level_summary.iterrows():
            label = " / ".join(str(row.get(c, "")).strip() for c in ["Ausweis_1", "Ausweis_2", "Ausweis_3"] if c in row and str(row.get(c, "")).strip())
            lines.append(
                f"- {label}: laufendes Jahr {format_de_amount(row['Laufendes Jahr'], 'EUR')}, "
                f"Vorjahr {format_de_amount(row['Vorjahr'], 'EUR')}, "
                f"Veränderung {format_de_amount(row['Veränderung'], 'EUR')} ({format_percent(row['Veränderung_%'])})."
            )

    lines.extend(["", "## Größte Kontentreiber"])
    if top_accounts.empty:
        lines.append("- Keine Konten oberhalb der Schwelle.")
    else:
        for _, row in top_accounts.iterrows():
            label = " / ".join(str(row.get(c, "")).strip() for c in account_cols if c not in ["KontoNr", "Kontobezeichnung"] and str(row.get(c, "")).strip())
            lines.append(
                f"- Konto {row.get('KontoNr', '')} {row.get('Kontobezeichnung', '')} ({label}): "
                f"laufendes Jahr {format_de_amount(row['Laufendes Jahr'], 'EUR')}, "
                f"Vorjahr {format_de_amount(row['Vorjahr'], 'EUR')}, "
                f"Veränderung {format_de_amount(row['Veränderung'], 'EUR')} ({format_percent(row['Veränderung_%'])})."
            )

    lines.extend(
        [
            "",
            "## Auftrag an die KI",
            "Bitte interpretiere die Entwicklung fachlich vorsichtig und prüferfreundlich.",
            f"Erstelle einen Text für: {ai_config.get('purpose', 'Management-Reporting, Anhang und Lagebericht')}.",
            f"Adressat: {ai_config.get('audience', 'Geschäftsführung und Wirtschaftsprüfer')}.",
            "",
            "## Wichtige Regeln",
            *(f"- {rule}" for rule in rules),
            "",
            "## Ausgabe-Struktur",
            *(f"{idx}. {item}" for idx, item in enumerate(output_structure, start=1)),
            "",
            "Wichtig: Keine Tatsachen erfinden. Formuliere Hypothesen als prüfbedürftig, wenn kein Sachverhalt geliefert wurde.",
        ]
    )
    return "\n".join(lines)


@st.cache_resource(show_spinner=False)
def get_openai_client():
    if OpenAI is None:
        return None
    try:
        api_key = st.secrets.get("OPENAI_API_KEY")
        if not api_key:
            return None
        return OpenAI(api_key=api_key)
    except Exception:
        return None


def openai_status() -> str:
    if OpenAI is None:
        return "OpenAI-Paket ist nicht installiert."
    if not st.secrets.get("OPENAI_API_KEY"):
        return "OPENAI_API_KEY fehlt in Streamlit Secrets."
    return "verbunden"


def generate_openai_interpretation(prompt_text: str, model: str, temperature: float, ai_config: dict) -> tuple[str | None, str | None]:
    client = get_openai_client()
    if client is None:
        return None, openai_status()

    instructions = f"""
Du bist ein erfahrener {ai_config.get('role', 'HGB-Abschlussanalyst')}.
Erzeuge eine fachlich belastbare, prüferfreundliche Interpretation der gelieferten Zahlenbasis.
Trenne klar zwischen beobachtbaren Zahlenentwicklungen, möglichen Ursachen/Hypothesen und Rückfragen.
Erfinde keine Sachverhalte. Nutze keine externen Informationen. Schreibe prägnant, aber verwertbar.
Zweck: {ai_config.get('purpose', 'Management-Reporting, Anhang-Hinweise und Lagebericht-Hinweise')}.
Adressat: {ai_config.get('audience', 'Geschäftsführung und Wirtschaftsprüfer')}.
Ton/Stil: {ai_config.get('tone', 'prüferfreundlich, vorsichtig, sachlich, konservativ')}.
Textumfang: {ai_config.get('length', 'mittel')}.
Form: {ai_config.get('form', 'gegliederte Abschnitte')}.
Formatierung:
- Verwende klare Überschriften ohne Markdown-Formatzeichen.
- Verwende keine Sternchen zur Hervorhebung.
- Verwende kurze Absätze und saubere Aufzählungen.
- Zahlen und Währungen bitte lesbar in deutscher Schreibweise wiedergeben.
"""

    try:
        kwargs = {"model": model, "instructions": instructions, "input": prompt_text}
        if model.startswith(("gpt-4", "gpt-4o")):
            kwargs["temperature"] = temperature
        response = client.responses.create(**kwargs)
        return clean_ai_output(response.output_text), None
    except Exception as e:
        return None, f"OpenAI-Anfrage fehlgeschlagen: {e}"


def _tree_amount_cells(row: pd.Series, value_cols: list[str]) -> str:
    return "".join(
        f"<span class='amount'>{format_de_number(row.get(c, 0))}</span>"
        for c in value_cols
    )


def _leaf_table_html(df: pd.DataFrame, value_cols: list[str]) -> str:
    header = "<tr><th>Konto-Nr.</th><th>Kontobezeichnung</th>"
    header += "".join(f"<th>{html.escape(str(c))}</th>" for c in value_cols)
    header += "</tr>"

    rows = []
    for _, row in df.sort_values("KontoNr").iterrows():
        cells = [
            html.escape(str(row.get("KontoNr", ""))),
            html.escape(str(row.get("Kontobezeichnung", ""))),
        ]
        cells.extend(format_de_number(row.get(c, 0)) for c in value_cols)
        rows.append("<tr>" + "".join(f"<td>{cell}</td>" for cell in cells) + "</tr>")

    return f"<table class='lumina-tree-table'>{header}{''.join(rows)}</table>"


def _tree_node_html(df: pd.DataFrame, levels: list[str], value_cols: list[str], depth: int = 0) -> str:
    if not levels:
        return _leaf_table_html(df, value_cols)

    col = levels[0]
    parts = []
    for label, group in df.groupby(col, dropna=False, sort=True):
        label_text = str(label).strip() if str(label).strip() else "(ohne weitere Untergliederung)"
        sums = group[value_cols].sum(numeric_only=True)
        summary = (
            f"<span class='tree-label'>{html.escape(label_text)}</span>"
            f"<span class='count'>{len(group)} Konto/Konten</span>"
            f"<span class='amounts'>{_tree_amount_cells(sums, value_cols)}</span>"
        )
        open_attr = " open" if depth < 1 else ""
        parts.append(
            f"<details class='lumina-tree-level level-{depth}'{open_attr}>"
            f"<summary>{summary}</summary>"
            f"{_tree_node_html(group, levels[1:], value_cols, depth + 1)}"
            f"</details>"
        )
    return "".join(parts)


def tree_view_html(df: pd.DataFrame, value_cols: list[str]) -> str:
    ausweis_cols = [f"Ausweis_{i}" for i in range(1, 8) if f"Ausweis_{i}" in df.columns]
    tree_df = df.copy()
    for col in ausweis_cols:
        tree_df[col] = tree_df[col].fillna("").astype(str).str.strip()
    tree_df = tree_df[tree_df["KontoNr"].astype(str).str.strip() != ""].copy()

    amount_headers = "".join(f"<span>{html.escape(str(c))}</span>" for c in value_cols)

    return f"""
    <style>
    .lumina-tree {{
        max-width: 100%;
    }}
    .lumina-tree-header {{
        display: grid;
        grid-template-columns: minmax(22rem, 1fr) 8rem repeat({len(value_cols)}, 10rem);
        column-gap: 0.75rem;
        padding: 0.35rem 0.3rem 0.45rem 2.3rem;
        border-bottom: 1px solid #e5e7eb;
        color: #6b7280;
        font-size: 0.88rem;
    }}
    .lumina-tree-header .amount-head {{
        display: contents;
    }}
    .lumina-tree-header span:not(:first-child) {{
        text-align: right;
    }}
    .lumina-tree details {{
        border-left: 1px solid #e5e7eb;
        margin: 0.18rem 0 0.18rem 0.9rem;
        padding-left: 0.7rem;
    }}
    .lumina-tree summary {{
        cursor: pointer;
        padding: 0.45rem 0.3rem;
        font-weight: 650;
        color: #1f2937;
        display: grid;
        grid-template-columns: minmax(22rem, 1fr) 8rem repeat({len(value_cols)}, 10rem);
        column-gap: 0.75rem;
        align-items: baseline;
    }}
    .lumina-tree summary .count {{
        color: #6b7280;
        font-weight: 400;
        font-size: 0.9rem;
        text-align: right;
        white-space: nowrap;
    }}
    .lumina-tree summary .amounts {{
        display: contents;
    }}
    .lumina-tree summary .amount {{
        text-align: right;
        font-variant-numeric: tabular-nums;
        white-space: nowrap;
        color: #4b5563;
        font-weight: 400;
    }}
    .lumina-tree-table {{
        border-collapse: collapse;
        width: 100%;
        margin: 0.25rem 0 0.8rem 1.2rem;
        font-size: 0.92rem;
    }}
    .lumina-tree-table th, .lumina-tree-table td {{
        border-bottom: 1px solid #e5e7eb;
        padding: 0.35rem 0.5rem;
        text-align: left;
    }}
    .lumina-tree-table td:nth-child(n+3), .lumina-tree-table th:nth-child(n+3) {{
        text-align: right;
        white-space: nowrap;
        font-variant-numeric: tabular-nums;
    }}
    </style>
    <div class="lumina-tree">
        <div class="lumina-tree-header">
            <span>Position</span>
            <span>Konten</span>
            <span class="amount-head">{amount_headers}</span>
        </div>
        {_tree_node_html(tree_df, ausweis_cols, value_cols)}
    </div>
    """


def _is_excel_value_column(header) -> bool:
    name = str(header).lower()
    if "konto" in name or "bezeichnung" in name or "text" in name:
        return False
    return any(x in name for x in ["saldo", "betrag", "wert", "summe", "202", "31.12", "haben", "soll", "debit", "credit", "balance"])


def _format_excel_numbers(workbook):
    number_format = '#,##0.00;[Red]-#,##0.00;0.00'
    for ws in workbook.worksheets:
        headers = [cell.value for cell in ws[1]]
        value_columns = {
            idx
            for idx, header in enumerate(headers, start=1)
            if _is_excel_value_column(header)
        }

        for row in ws.iter_rows(min_row=2):
            for cell in row:
                if cell.column in value_columns and isinstance(cell.value, (int, float)):
                    cell.number_format = number_format
                    cell.alignment = cell.alignment.copy(horizontal="right")


def excel_export(susa_raw: pd.DataFrame, mapped: pd.DataFrame, pivot: pd.DataFrame, klarung: pd.DataFrame) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        susa_raw.to_excel(writer, sheet_name="01_Rohdaten", index=False)
        mapped.to_excel(writer, sheet_name="02_Mapping_Ergebnis", index=False)
        klarung.to_excel(writer, sheet_name="03_Klaerungsposten", index=False)
        pivot.to_excel(writer, sheet_name="04_Bilanz_GuV", index=False)

        workbook = writer.book
        _format_excel_numbers(workbook)
        for ws in workbook.worksheets:
            ws.freeze_panes = "A2"
            for col_cells in ws.columns:
                max_len = max(len(str(cell.value)) if cell.value is not None else 0 for cell in col_cells)
                ws.column_dimensions[col_cells[0].column_letter].width = min(max(max_len + 2, 10), 45)

    return buffer.getvalue()


def template_bytes(filename: str) -> bytes | None:
    path = TEMPLATE_DIR / filename
    if not path.exists():
        return None
    return path.read_bytes()


# ------------------------------------------------------------
# Supabase optional
# ------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def get_supabase_client():
    if create_client is None:
        return None
    try:
        url = st.secrets.get("SUPABASE_URL")
        key = st.secrets.get("SUPABASE_KEY")
        if not url or not key:
            return None
        return create_client(url, key)
    except Exception:
        return None


def list_mapping_names_from_supabase() -> tuple[list[str], str | None]:
    sb = get_supabase_client()
    if sb is None:
        return [DEFAULT_MAPPING_NAME], "Supabase ist nicht verbunden."
    try:
        res = sb.table("master_mapping").select("mapping_name").execute()
        names = sorted({
            str(row.get("mapping_name") or DEFAULT_MAPPING_NAME).strip()
            for row in (res.data or [])
            if str(row.get("mapping_name") or DEFAULT_MAPPING_NAME).strip()
        })
        return names or [DEFAULT_MAPPING_NAME], None
    except Exception:
        return [DEFAULT_MAPPING_NAME], "Mehrere Master-Mappings benötigen in Supabase die Spalte 'mapping_name'. Bis dahin wird 'Standard' verwendet."


def load_mapping_from_supabase(mapping_name: str = DEFAULT_MAPPING_NAME):
    sb = get_supabase_client()
    if sb is None:
        return None, "Supabase ist nicht verbunden. Prüfe Streamlit Secrets."
    try:
        rows = []
        page_size = 1000
        start = 0
        supports_mapping_name = True

        while True:
            end = start + page_size - 1
            query = sb.table("master_mapping").select("*").order("konto_nr").range(start, end)
            if mapping_name != DEFAULT_MAPPING_NAME:
                query = query.eq("mapping_name", mapping_name)
            elif supports_mapping_name:
                query = query.or_(f"mapping_name.eq.{DEFAULT_MAPPING_NAME},mapping_name.is.null")
            try:
                res = query.execute()
            except Exception:
                if mapping_name != DEFAULT_MAPPING_NAME:
                    return None, "Dieses Mapping konnte nicht geladen werden. Prüfe, ob die Supabase-Spalte 'mapping_name' existiert."
                supports_mapping_name = False
                res = (
                    sb.table("master_mapping")
                    .select("*")
                    .order("konto_nr")
                    .range(start, end)
                    .execute()
                )
            batch = res.data or []
            rows.extend(batch)

            if len(batch) < page_size:
                break
            start += page_size

        df = pd.DataFrame(rows)
        if df.empty:
            return None, "Tabelle master_mapping ist leer."
        df = df.rename(columns={"konto_nr": "KontoNr"})
        for i in range(1, 8):
            df = df.rename(columns={f"ausweis_{i}": f"Ausweis_{i}"})
        return normalize_mapping(df), None
    except Exception as e:
        return None, f"Supabase-Laden fehlgeschlagen: {e}"


def save_mapping_to_supabase(mapping: pd.DataFrame, mapping_name: str = DEFAULT_MAPPING_NAME):
    sb = get_supabase_client()
    if sb is None:
        return "Supabase ist nicht verbunden."
    try:
        df = normalize_mapping(mapping)
        rows = []
        for _, r in df.iterrows():
            item = {"mapping_name": mapping_name.strip() or DEFAULT_MAPPING_NAME, "konto_nr": r["KontoNr"]}
            for i in range(1, 8):
                item[f"ausweis_{i}"] = r[f"Ausweis_{i}"]
            rows.append(item)
        try:
            sb.table("master_mapping").upsert(rows, on_conflict="mapping_name,konto_nr").execute()
        except Exception as e:
            if mapping_name.strip() and mapping_name.strip() != DEFAULT_MAPPING_NAME:
                return (
                    "Supabase-Speichern fehlgeschlagen: Für mehrere Master-Mappings muss in Supabase "
                    "die Migration 'supabase_master_mapping_migration.sql' ausgeführt werden. "
                    "Aktuell ist vermutlich noch 'konto_nr' alleiniger Primärschlüssel. "
                    f"Details: {e}"
                )
            legacy_rows = [{k: v for k, v in row.items() if k != "mapping_name"} for row in rows]
            sb.table("master_mapping").upsert(legacy_rows, on_conflict="konto_nr").execute()
        return None
    except Exception as e:
        return f"Supabase-Speichern fehlgeschlagen: {e}"


def new_id() -> str:
    return str(uuid.uuid4())


def stable_id(*parts) -> str:
    return str(uuid.uuid5(uuid.NAMESPACE_URL, "|".join(str(p or "") for p in parts)))


def now_iso() -> str:
    return datetime.utcnow().isoformat()


def table_missing_message(table_name: str) -> str:
    return f"Supabase-Tabelle '{table_name}' ist nicht verfügbar. Bitte SQL-Migration ausführen."


def supabase_error_message(table_name: str, error: Exception) -> str:
    text = str(error)
    if "row-level security" in text.lower() or "42501" in text:
        return f"Supabase-RLS blockiert '{table_name}'. Bitte RLS-Policies für App-Zugriff anlegen oder Service-Key verwenden. Details: {error}"
    return f"{table_missing_message(table_name)} Details: {error}"


def supabase_select(table_name: str, filters: dict | None = None, order: str | None = None) -> tuple[list[dict], str | None]:
    sb = get_supabase_client()
    if sb is None:
        return [], "Supabase ist nicht verbunden."
    try:
        query = sb.table(table_name).select("*")
        for key, value in (filters or {}).items():
            query = query.eq(key, value)
        if order:
            query = query.order(order)
        res = query.execute()
        return res.data or [], None
    except Exception as e:
        return [], supabase_error_message(table_name, e)


def supabase_upsert(table_name: str, rows, on_conflict: str = "id") -> str | None:
    sb = get_supabase_client()
    if sb is None:
        return "Supabase ist nicht verbunden."
    try:
        sb.table(table_name).upsert(rows, on_conflict=on_conflict).execute()
        return None
    except Exception as e:
        return supabase_error_message(table_name, e)


def supabase_delete(table_name: str, row_id: str) -> str | None:
    sb = get_supabase_client()
    if sb is None:
        return "Supabase ist nicht verbunden."
    try:
        sb.table(table_name).delete().eq("id", row_id).execute()
        return None
    except Exception as e:
        return supabase_error_message(table_name, e)


def audit_log(action: str, description: str, mandant_id: str | None = None, year_id: str | None = None):
    row = {
        "id": new_id(),
        "mandant_id": mandant_id or st.session_state.get("active_mandant_id"),
        "year_id": year_id or st.session_state.get("active_year_id"),
        "action": action,
        "description": description,
        "timestamp": now_iso(),
        "user": st.secrets.get("APP_USER", "streamlit-user") if hasattr(st, "secrets") else "streamlit-user",
    }
    supabase_upsert("audit_log", row)


def load_mandants() -> tuple[list[dict], str | None]:
    return supabase_select("mandants", order="mandantenname")


def load_years(mandant_id: str | None) -> tuple[list[dict], str | None]:
    if not mandant_id:
        return [], None
    return supabase_select("mandant_years", {"mandant_id": mandant_id}, order="jahr")


def load_reporting_profile(mandant_id: str | None) -> tuple[dict, str | None]:
    if not mandant_id:
        return {}, None
    rows, err = supabase_select("reporting_profiles", {"mandant_id": mandant_id})
    return (rows[0] if rows else {}, err)


def load_onboarding_answers(mandant_id: str | None, year_id: str | None) -> tuple[dict, str | None]:
    if not mandant_id:
        return {}, None
    filters = {"mandant_id": mandant_id}
    if year_id:
        filters["year_id"] = year_id
    rows, err = supabase_select("onboarding_answers", filters)
    answers = {(r.get("section"), r.get("question")): r.get("answer", "") for r in rows}
    return answers, err


def load_susa_uploads(mandant_id: str | None, year_id: str | None) -> tuple[list[dict], str | None]:
    if not mandant_id or not year_id:
        return [], None
    return supabase_select("susa_uploads", {"mandant_id": mandant_id, "year_id": year_id}, order="created_at")


def combined_susa_frame(susa_files: list[dict]) -> tuple[pd.DataFrame | None, dict]:
    frames = [item["norm"] for item in susa_files if isinstance(item.get("norm"), pd.DataFrame)]
    if not frames:
        return None, {}
    combined = pd.concat(frames, ignore_index=True)
    value_cols = detect_value_cols(combined)
    group_cols = ["KontoNr", "Kontobezeichnung"]
    other_cols = [c for c in combined.columns if c not in group_cols + value_cols]
    agg = combined.groupby(group_cols, dropna=False)[value_cols].sum().reset_index()
    for col in other_cols:
        if col not in agg.columns and col in combined.columns:
            first_values = combined.groupby("KontoNr", dropna=False)[col].first().reset_index()
            agg = agg.merge(first_values, on="KontoNr", how="left")
    return agg, {"value_cols": value_cols, "source_count": len(frames)}


# ------------------------------------------------------------
# Session defaults
# ------------------------------------------------------------
for key, default in {
    "mandant": "Beispiel GmbH",
    "abschlussjahr": datetime.now().year - 1,
    "standard": "HGB Einzelabschluss",
    "active_mandant_id": None,
    "active_year_id": None,
    "active_mandant": {},
    "active_year": {},
    "reporting_profile": {},
    "onboarding_answers": {},
    "mapping_name": DEFAULT_MAPPING_NAME,
    "susa_raw": None,
    "susa_norm": None,
    "susa_files": [],
    "selected_susa_scope": "Gesamt",
    "mapping": None,
    "mapped": None,
    "meta": {},
    "flash_message": None,
    "ai_interpretation": "",
    "export_status": "nicht erzeugt",
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


# ------------------------------------------------------------
# Sidebar
# ------------------------------------------------------------
st.sidebar.title("LUMINA")
st.sidebar.caption(f"Version {APP_VERSION}")

phase = st.sidebar.radio(
    "Navigation",
    [
        "1 Willkommen",
        "2 Mandanten",
        "3 Onboarding",
        "4 Upload & Mapping",
        "5 Prüfen",
        "6 Abschlussansicht",
        "7 Interpretation",
        "8 Export",
    ],
)

sb_client = get_supabase_client()
clarification_count = mapping_counts(st.session_state.mapped)["klarung"] if st.session_state.mapped is not None else 0
st.sidebar.markdown("---")
st.sidebar.write("**Status**")
st.sidebar.write("Supabase:", "✅ verbunden" if sb_client else "⚠️ nicht verbunden")
st.sidebar.write("OpenAI:", "✅ verbunden" if openai_status() == "verbunden" else "⚠️ nicht verbunden")
st.sidebar.write("Aktiver Mandant:", st.session_state.active_mandant.get("mandantenname") or st.session_state.mandant)
st.sidebar.write("Aktives Jahr:", st.session_state.active_year.get("jahr") or st.session_state.abschlussjahr)
st.sidebar.write("Anzahl SuSAs:", len(st.session_state.susa_files))
st.sidebar.write("Master-Mapping:", st.session_state.mapping_name)
st.sidebar.write("Mapping:", "✅ geladen" if st.session_state.mapping is not None else "⚠️ fehlt")
st.sidebar.write("SuSa:", "✅ geladen" if st.session_state.susa_norm is not None else "⚠️ fehlt")
st.sidebar.write("Klärung:", clarification_count)
st.sidebar.write("Export:", st.session_state.export_status)


# ------------------------------------------------------------
# Header
# ------------------------------------------------------------
st.title("📊 LUMINA Mapping – internes Abschluss-Cockpit")
st.caption("Von der SuSa zum prüferfreundlichen Excel-Abschluss-Paket.")


# ------------------------------------------------------------
# Phase 1
# ------------------------------------------------------------
if phase == "1 Willkommen":
    st.markdown(
        """
        <div class="lumina-box">
        <b>Ziel dieser App:</b><br>
        Du lädst eine Mandanten-SuSa und ein Master-Mapping hoch. LUMINA erzeugt daraus eine strukturierte HGB-Bilanz/GuV,
        zeigt Klärungsposten und erstellt ein Excel-Prüfungspaket.
        </div>
        """,
        unsafe_allow_html=True,
    )

    c1, c2, c3 = st.columns(3)
    c1.metric("1", "SuSa hochladen")
    c2.metric("2", "Mapping prüfen")
    c3.metric("3", "Excel exportieren")

    st.info("Empfehlung: Diese App zuerst als internes Werkzeug nutzen. Kundendaten separat über Smartdocu/SharePoint austauschen.")


# ------------------------------------------------------------
# Phase 2
# ------------------------------------------------------------
elif phase == "2 Mandanten":
    st.subheader("Mandanten & Abschlussjahre")

    mandants, mandant_err = load_mandants()
    if mandant_err:
        st.warning(mandant_err)

    if mandants:
        labels = {f"{m.get('mandantenname', 'Ohne Namen')} ({m.get('rechtsform', '-')})": m for m in mandants}
        current_label = next((label for label, item in labels.items() if item.get("id") == st.session_state.active_mandant_id), list(labels.keys())[0])
        selected_label = st.selectbox("Bestehenden Mandanten auswählen", list(labels.keys()), index=list(labels.keys()).index(current_label))
        selected_mandant = labels[selected_label]
        if st.button("Mandant aktiv setzen", use_container_width=True):
            st.session_state.active_mandant = selected_mandant
            st.session_state.active_mandant_id = selected_mandant.get("id")
            st.session_state.mandant = selected_mandant.get("mandantenname", st.session_state.mandant)
            st.success(f"Aktiver Mandant: {st.session_state.mandant}")
            audit_log("select_mandant", f"Mandant aktiviert: {st.session_state.mandant}", st.session_state.active_mandant_id)
    else:
        st.info("Noch kein Mandant aus Supabase geladen. Du kannst unten einen Mandanten anlegen.")

    st.markdown("### Mandanten-Stammdaten")
    current = st.session_state.active_mandant or {}
    with st.form("mandant_form"):
        col1, col2, col3 = st.columns(3)
        with col1:
            mandantenname = st.text_input("Mandantenname", current.get("mandantenname", st.session_state.mandant))
            rechtsform = st.text_input("Rechtsform", current.get("rechtsform", ""))
            branche = st.text_input("Branche", current.get("branche", ""))
            sitz = st.text_input("Sitz", current.get("sitz", ""))
        with col2:
            geschaeftsjahr = st.text_input("Geschäftsjahr", current.get("geschaeftsjahr", "01.01.-31.12."))
            kontenrahmen = st.text_input("Kontenrahmen", current.get("kontenrahmen", st.session_state.mapping_name))
            groessenklasse = st.selectbox("Größenklasse", ["kleinst", "klein", "mittelgroß", "groß"], index=1)
            pruefungspflicht = st.checkbox("Prüfungspflicht", value=bool(current.get("pruefungspflicht", False)))
            lageberichtspflicht = st.checkbox("Lageberichtspflicht", value=bool(current.get("lageberichtspflicht", False)))
        with col3:
            steuerberater = st.text_input("Steuerberater", current.get("steuerberater", ""))
            wirtschaftspruefer = st.text_input("Wirtschaftsprüfer", current.get("wirtschaftspruefer", ""))
            ansprechpartner = st.text_input("Ansprechpartner Rechnungswesen", current.get("ansprechpartner_rechnungswesen", ""))
            besonderheiten = st.text_area("Besonderheiten", current.get("besonderheiten", ""), height=120)

        save_mandant = st.form_submit_button("Mandant speichern", type="primary", use_container_width=True)

    if save_mandant:
        row = {
            "id": current.get("id") or new_id(),
            "mandantenname": mandantenname,
            "rechtsform": rechtsform,
            "branche": branche,
            "sitz": sitz,
            "geschaeftsjahr": geschaeftsjahr,
            "kontenrahmen": kontenrahmen,
            "groessenklasse": groessenklasse,
            "pruefungspflicht": pruefungspflicht,
            "lageberichtspflicht": lageberichtspflicht,
            "steuerberater": steuerberater,
            "wirtschaftspruefer": wirtschaftspruefer,
            "ansprechpartner_rechnungswesen": ansprechpartner,
            "besonderheiten": besonderheiten,
            "updated_at": now_iso(),
        }
        if "created_at" not in current:
            row["created_at"] = now_iso()
        err = supabase_upsert("mandants", row)
        if err:
            st.error(err)
        else:
            st.session_state.active_mandant = row
            st.session_state.active_mandant_id = row["id"]
            st.session_state.mandant = row["mandantenname"]
            st.success("Mandant gespeichert.")
            audit_log("save_mandant", f"Mandant gespeichert: {row['mandantenname']}", row["id"])

    if st.session_state.active_mandant_id:
        st.markdown("### Abschlussjahre")
        years, year_err = load_years(st.session_state.active_mandant_id)
        if year_err:
            st.warning(year_err)
        if years:
            year_labels = {str(y.get("jahr")): y for y in years}
            selected_year_label = st.selectbox("Abschlussjahr auswählen", list(year_labels.keys()))
            if st.button("Jahr aktiv setzen", use_container_width=True):
                year = year_labels[selected_year_label]
                st.session_state.active_year = year
                st.session_state.active_year_id = year.get("id")
                st.session_state.abschlussjahr = int(year.get("jahr"))
                st.success(f"Aktives Jahr: {st.session_state.abschlussjahr}")

        with st.form("year_form"):
            col_y1, col_y2, col_y3 = st.columns(3)
            with col_y1:
                jahr = st.number_input("Neues/zu speicherndes Jahr", min_value=2000, max_value=2100, value=int(st.session_state.abschlussjahr))
            with col_y2:
                status = st.selectbox("Status", ["angelegt", "in Bearbeitung", "klärungsbedürftig", "abgeschlossen"], index=1)
                materiality = st.number_input("Wesentlichkeitsschwelle", min_value=0.0, value=10000.0, step=1000.0)
            with col_y3:
                wesentliche_themen = st.text_area("Wesentliche Themen", height=100)
            save_year = st.form_submit_button("Abschlussjahr speichern", use_container_width=True)

        if save_year:
            existing = next((y for y in years if int(y.get("jahr")) == int(jahr)), {})
            row = {
                "id": existing.get("id") or new_id(),
                "mandant_id": st.session_state.active_mandant_id,
                "jahr": int(jahr),
                "status": status,
                "wesentliche_themen": wesentliche_themen,
                "materiality_threshold": materiality,
                "updated_at": now_iso(),
            }
            if "created_at" not in existing:
                row["created_at"] = now_iso()
            err = supabase_upsert("mandant_years", row)
            if err:
                st.error(err)
            else:
                st.session_state.active_year = row
                st.session_state.active_year_id = row["id"]
                st.session_state.abschlussjahr = int(jahr)
                st.success("Abschlussjahr gespeichert.")
                audit_log("save_year", f"Abschlussjahr gespeichert: {jahr}", st.session_state.active_mandant_id, row["id"])

        st.markdown("### Löschen")
        delete_confirm = st.text_input("Zum Löschen des aktiven Mandanten bitte Mandantenname exakt eingeben")
        if st.button("Aktiven Mandanten löschen", disabled=delete_confirm != st.session_state.active_mandant.get("mandantenname")):
            err = supabase_delete("mandants", st.session_state.active_mandant_id)
            if err:
                st.error(err)
            else:
                audit_log("delete_mandant", f"Mandant gelöscht: {st.session_state.active_mandant.get('mandantenname')}", st.session_state.active_mandant_id)
                st.session_state.active_mandant = {}
                st.session_state.active_mandant_id = None
                st.success("Mandant gelöscht.")


# ------------------------------------------------------------
# Phase 3
# ------------------------------------------------------------
elif phase == "3 Onboarding":
    st.subheader("Mandanten-Onboarding / Smart Interview")

    if not st.session_state.active_mandant_id:
        st.warning("Bitte zuerst unter 'Mandanten' einen Mandanten aktiv setzen.")
    else:
        answers, ans_err = load_onboarding_answers(st.session_state.active_mandant_id, st.session_state.active_year_id)
        if ans_err:
            st.warning(ans_err)
        st.session_state.onboarding_answers = answers

        profile, profile_err = load_reporting_profile(st.session_state.active_mandant_id)
        if profile_err:
            st.warning(profile_err)
        st.session_state.reporting_profile = profile

        st.markdown("### Reporting-Profil")
        with st.form("reporting_profile_form"):
            c1, c2, c3 = st.columns(3)
            with c1:
                berichtsstil = st.selectbox("Berichtsstil", ["konservativ", "neutral", "managementorientiert"], index=0)
                textumfang = st.selectbox("Textumfang", ["kurz", "mittel", "lang"], index=1)
            with c2:
                ausgabeform = st.selectbox("Ausgabeform", ["Fließtext", "Bulletpoints", "Tabelle"], index=1)
                zielgruppe = st.selectbox("Zielgruppe", ["Geschäftsführung", "Bank", "Wirtschaftsprüfer", "internes Rechnungswesen"], index=0)
            with c3:
                anhang_level = st.selectbox("Anhang-Level", ["minimal", "erweitert", "prüferfreundlich"], index=2)
                lagebericht_stil = st.selectbox("Lagebericht-Stil", ["sachlich", "strategisch", "vorsichtig"], index=2)
            save_profile = st.form_submit_button("Reporting-Profil speichern", type="primary", use_container_width=True)

        if save_profile:
            row = {
                "id": profile.get("id") or stable_id(st.session_state.active_mandant_id, "reporting_profile"),
                "mandant_id": st.session_state.active_mandant_id,
                "berichtsstil": berichtsstil,
                "textumfang": textumfang,
                "ausgabeform": ausgabeform,
                "zielgruppe": zielgruppe,
                "anhang_level": anhang_level,
                "lagebericht_stil": lagebericht_stil,
                "updated_at": now_iso(),
            }
            err = supabase_upsert("reporting_profiles", row)
            if err:
                st.error(err)
            else:
                st.session_state.reporting_profile = row
                st.success("Reporting-Profil gespeichert.")
                audit_log("save_reporting_profile", "Reporting-Profil gespeichert")

        st.markdown("### Smart Interview")
        with st.form("onboarding_form"):
            answer_rows = []
            for section, questions in ONBOARDING_SECTIONS.items():
                st.markdown(f"#### {section}")
                for question in questions:
                    key = (section, question)
                    answer = st.text_area(question, value=answers.get(key, ""), key=f"onb_{section}_{question}", height=90)
                    answer_rows.append(
                        {
                            "id": stable_id(st.session_state.active_mandant_id, st.session_state.active_year_id, section, question),
                            "mandant_id": st.session_state.active_mandant_id,
                            "year_id": st.session_state.active_year_id,
                            "section": section,
                            "question": question,
                            "answer": answer,
                            "updated_at": now_iso(),
                        }
                    )
            save_answers = st.form_submit_button("Onboarding speichern", use_container_width=True)

        if save_answers:
            rows = [row for row in answer_rows if row["answer"].strip()]
            err = supabase_upsert("onboarding_answers", rows) if rows else None
            if err:
                st.error(err)
            else:
                st.success("Onboarding-Antworten gespeichert.")
                audit_log("save_onboarding", f"{len(rows)} Onboarding-Antworten gespeichert")


# Phase 4
# ------------------------------------------------------------
elif phase == "4 Upload & Mapping":
    st.subheader("Upload & Mapping")

    left, right = st.columns(2)

    with left:
        st.markdown("#### 1. Master-Mapping")
        mapping_names, mapping_names_warning = list_mapping_names_from_supabase()
        if st.session_state.mapping_name not in mapping_names:
            mapping_names = [st.session_state.mapping_name] + mapping_names

        selected_mapping_name = st.selectbox(
            "Master-Mapping auswählen",
            mapping_names,
            index=mapping_names.index(st.session_state.mapping_name),
        )
        new_mapping_name = st.text_input("Name für neues/zu speicherndes Master-Mapping", selected_mapping_name)
        st.session_state.mapping_name = (new_mapping_name or DEFAULT_MAPPING_NAME).strip()
        if mapping_names_warning:
            st.caption(mapping_names_warning)

        mapping_file = st.file_uploader("Mapping-Excel hochladen", type=["xlsx", "xls"], key="mapping_file")

        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("☁️ Mapping aus Supabase laden", use_container_width=True):
                df, err = load_mapping_from_supabase(st.session_state.mapping_name)
                if err:
                    st.warning(err)
                else:
                    st.session_state.mapping = df
                    st.success(f"{len(df)} Mapping-Zeilen für '{st.session_state.mapping_name}' geladen.")

        with col_b:
            if st.button("💾 Mapping nach Supabase sichern", use_container_width=True):
                if st.session_state.mapping is None:
                    st.warning("Noch kein Mapping geladen.")
                else:
                    err = save_mapping_to_supabase(st.session_state.mapping, st.session_state.mapping_name)
                    if err:
                        st.error(err)
                    else:
                        st.success(f"Mapping '{st.session_state.mapping_name}' gespeichert.")
                        audit_log("save_mapping", f"Master-Mapping gespeichert: {st.session_state.mapping_name}")

        if mapping_file is not None:
            try:
                raw_map = read_excel_smart(mapping_file)
                st.session_state.mapping = normalize_mapping(raw_map)
                st.success(f"Mapping geladen: {len(st.session_state.mapping)} Konten. Speichern legt es unter '{st.session_state.mapping_name}' ab.")
            except Exception as e:
                st.error(str(e))

        if st.session_state.mapping is not None:
            st.dataframe(st.session_state.mapping.head(20), use_container_width=True, hide_index=True)

    with right:
        st.markdown("#### 2. Mandanten-SuSa")
        susa_type = st.selectbox("SuSa-Typ", SUSA_TYPES, index=0)
        susa_version = st.text_input("SuSa-Version", value=f"v{len(st.session_state.susa_files) + 1}")
        susa_files = st.file_uploader("SuSa-Excel hochladen", type=["xlsx", "xls"], key="susa_file", accept_multiple_files=True)

        st.markdown("#### Beispiele")
        sample_mapping = template_bytes("Muster_Kontenmapping.xlsx")
        sample_susa = template_bytes("Muster_Susa.xlsx")
        sample_col_1, sample_col_2 = st.columns(2)
        with sample_col_1:
            if sample_mapping:
                if st.button("Muster-Mapping laden", use_container_width=True):
                    try:
                        raw_map = read_excel_smart(io.BytesIO(sample_mapping))
                        st.session_state.mapping = normalize_mapping(raw_map)
                        st.success(f"Muster-Mapping geladen: {len(st.session_state.mapping)} Konten.")
                    except Exception as e:
                        st.error(str(e))
                st.download_button(
                    "Muster-Mapping herunterladen",
                    data=sample_mapping,
                    file_name="Muster_Kontenmapping.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
        with sample_col_2:
            if sample_susa:
                if st.button("Muster-SuSa laden", use_container_width=True):
                    try:
                        raw = read_excel_smart(io.BytesIO(sample_susa))
                        norm, meta = normalize_susa(raw)
                        st.session_state.susa_raw = raw
                        st.session_state.susa_norm = norm
                        st.session_state.meta = meta
                        st.success(f"Muster-SuSa geladen: {len(norm)} Konten.")
                    except Exception as e:
                        st.error(str(e))
                st.download_button(
                    "Muster-SuSa herunterladen",
                    data=sample_susa,
                    file_name="Muster_Susa.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )

        if susa_files:
            for susa_file in susa_files:
                already_loaded = any(item.get("filename") == susa_file.name and item.get("version") == susa_version for item in st.session_state.susa_files)
                if already_loaded:
                    continue
                try:
                    raw = read_excel_smart(susa_file)
                    norm, meta = normalize_susa(raw)
                    item = {
                        "id": new_id(),
                        "name": susa_file.name,
                        "filename": susa_file.name,
                        "type": susa_type,
                        "version": susa_version,
                        "raw": raw,
                        "norm": norm,
                        "meta": meta,
                        "created_at": now_iso(),
                    }
                    st.session_state.susa_files.append(item)
                    st.session_state.susa_raw = raw
                    st.session_state.susa_norm = norm
                    st.session_state.meta = meta
                    if st.session_state.active_mandant_id and st.session_state.active_year_id:
                        supabase_upsert(
                            "susa_uploads",
                            {
                                "id": item["id"],
                                "mandant_id": st.session_state.active_mandant_id,
                                "year_id": st.session_state.active_year_id,
                                "name": item["name"],
                                "susa_type": item["type"],
                                "version": item["version"],
                                "row_count": len(norm),
                                "created_at": item["created_at"],
                            },
                        )
                        audit_log("upload_susa", f"SuSa hochgeladen: {item['name']} ({item['type']}, {item['version']})")
                    st.success(f"SuSa geladen: {susa_file.name} mit {len(norm)} Konten.")
                except Exception as e:
                    st.error(f"{susa_file.name}: {e}")

        if st.session_state.susa_files:
            st.markdown("#### Geladene SuSAs")
            overview = pd.DataFrame(
                [
                    {"Name": item["name"], "Typ": item["type"], "Version": item["version"], "Konten": len(item["norm"])}
                    for item in st.session_state.susa_files
                ]
            )
            st.dataframe(overview, use_container_width=True, hide_index=True)

            scope_options = ["Gesamt"] + [f"{item['name']} | {item['version']}" for item in st.session_state.susa_files]
            st.session_state.selected_susa_scope = st.selectbox("Auswertung", scope_options, index=0)
            if st.session_state.selected_susa_scope == "Gesamt":
                combined, combined_meta = combined_susa_frame(st.session_state.susa_files)
                if combined is not None:
                    st.session_state.susa_norm = combined
                    st.session_state.susa_raw = combined
                    st.session_state.meta = {"value_cols": combined_meta.get("value_cols", [])}
                    st.success(f"Gruppen-SuSa aggregiert: {len(combined)} Konten aus {combined_meta.get('source_count', 0)} SuSa-Dateien.")
                    st.dataframe(display_dataframe(combined.head(20), st.session_state.meta["value_cols"]), use_container_width=True, hide_index=True)
            else:
                selected = st.session_state.susa_files[scope_options.index(st.session_state.selected_susa_scope) - 1]
                st.session_state.susa_norm = selected["norm"]
                st.session_state.susa_raw = selected["raw"]
                st.session_state.meta = selected["meta"]
                st.success(f"Auswertung für einzelne SuSa: {selected['name']}")
                st.dataframe(display_dataframe(selected["norm"].head(20), selected["meta"]["value_cols"]), use_container_width=True, hide_index=True)

    st.markdown("---")
    if st.button("🚀 Mapping starten", type="primary", use_container_width=True):
        if st.session_state.susa_norm is None:
            st.error("Bitte zuerst eine SuSa hochladen.")
        elif st.session_state.mapping is None and get_dynamic_mapping is None:
            st.error("Bitte zuerst ein Mapping hochladen oder aus Supabase laden.")
        else:
            active_mapping = st.session_state.mapping if st.session_state.mapping is not None else empty_mapping_frame()
            st.session_state.mapped = apply_mapping(st.session_state.susa_norm, active_mapping)
            counts = mapping_counts(st.session_state.mapped)
            st.success(f"Mapping abgeschlossen. Vorschläge: {counts['vorschlag']}. Klärungsposten: {counts['klarung']}")
            st.dataframe(display_dataframe(st.session_state.mapped.head(50), st.session_state.meta.get("value_cols", [])), use_container_width=True, hide_index=True)


# ------------------------------------------------------------
# Phase 4
# ------------------------------------------------------------
elif phase == "5 Prüfen":
    st.subheader("Prüfen & Optimieren")
    if st.session_state.flash_message:
        st.success(st.session_state.flash_message)
        st.session_state.flash_message = None

    if st.session_state.mapped is None:
        st.warning("Bitte zuerst in Phase 4 das Mapping starten.")
    else:
        df = st.session_state.mapped.copy()
        value_cols = st.session_state.meta.get("value_cols", detect_value_cols(df))
        klarung = df[clarification_mask(df)].copy()
        counts = mapping_counts(df)

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Konten gesamt", f"{counts['total']:,}".replace(",", "."))
        c2.metric("Gemappt", f"{counts['gemappt']:,}".replace(",", "."))
        c3.metric("Vorschläge", f"{counts['vorschlag']:,}".replace(",", "."))
        c4.metric("Klärung", f"{counts['klarung']:,}".replace(",", "."))

        if value_cols and not klarung.empty:
            main_value = value_cols[-1]
            st.metric("Summe Klärungsposten", format_de_amount(klarung[main_value].sum(), "EUR"))

        if klarung.empty:
            st.success("Keine Klärungsposten gefunden.")
        else:
            st.error("Es gibt noch ungeklärte Konten. Diese solltest du im Master-Mapping ergänzen.")
            show_cols = ["KontoNr", "Kontobezeichnung"] + value_cols + [f"Ausweis_{i}" for i in range(1, 8)]
            show_cols = [c for c in show_cols if c in klarung.columns]
            st.dataframe(display_dataframe(klarung[show_cols], value_cols), use_container_width=True, hide_index=True)

            st.markdown("### Klärungsposten zuordnen")
            editor_cols = ["KontoNr", "Kontobezeichnung"] + value_cols + [f"Ausweis_{i}" for i in range(1, 8)]
            editor_cols = [c for c in editor_cols if c in klarung.columns]
            editor_df = klarung[editor_cols].copy()
            for i in range(1, 8):
                col = f"Ausweis_{i}"
                if col in editor_df.columns:
                    editor_df[col] = editor_df[col].apply(clean_manual_mapping_value)

            edited = st.data_editor(
                editor_df,
                key="klarung_editor",
                use_container_width=True,
                hide_index=True,
                num_rows="fixed",
                disabled=["KontoNr", "Kontobezeichnung"] + [c for c in value_cols if c in editor_df.columns],
                column_config=ausweis_column_config(st.session_state.mapping, editor_df),
            )

            save_col, download_col = st.columns(2)
            with save_col:
                if st.button("Zuordnungen ins Master-Mapping übernehmen", type="primary", use_container_width=True):
                    updates, skipped = rows_to_mapping_updates(edited)
                    if updates.empty:
                        st.warning("Bitte mindestens Ausweis_1 für eine Klärungszeile ausfüllen.")
                    else:
                        st.session_state.mapping = merge_mapping_updates(st.session_state.mapping, updates)
                        st.session_state.mapped = apply_mapping(st.session_state.susa_norm, st.session_state.mapping)
                        err = save_mapping_to_supabase(st.session_state.mapping, st.session_state.mapping_name)
                        if err:
                            st.error(err)
                        else:
                            message = f"{len(updates)} Zuordnung(en) ins Master-Mapping '{st.session_state.mapping_name}' übernommen und nach Supabase gespeichert."
                            if skipped:
                                message += f" Nicht übernommen, weil Ausweis_1 fehlt: {', '.join(skipped)}"
                            st.session_state.flash_message = message
                            audit_log("update_mapping", message)
                            st.rerun()

            with download_col:
                if st.session_state.mapping is not None:
                    mapping_buffer = io.BytesIO()
                    with pd.ExcelWriter(mapping_buffer, engine="openpyxl") as writer:
                        st.session_state.mapping.to_excel(writer, sheet_name="Master_Mapping", index=False)
                    st.download_button(
                        "Aktuelles Master-Mapping herunterladen",
                        data=mapping_buffer.getvalue(),
                        file_name="Lumina_Master_Mapping_aktualisiert.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                    )

        with st.expander("Vollständiges Mapping-Ergebnis anzeigen"):
            st.dataframe(display_dataframe(df, value_cols), use_container_width=True, hide_index=True)


# ------------------------------------------------------------
# Phase 5
# ------------------------------------------------------------
elif phase == "6 Abschlussansicht":
    st.subheader("Abschlussansicht")

    if st.session_state.mapped is None:
        st.warning("Bitte zuerst in Phase 4 das Mapping starten.")
    else:
        df = st.session_state.mapped.copy()
        value_cols = st.session_state.meta.get("value_cols", detect_value_cols(df))

        tree_tab, pivot_tab = st.tabs(["Baumansicht", "Pivot-Tabelle"])
        with tree_tab:
            if not value_cols:
                st.error("Keine Wertspalten gefunden.")
            else:
                st.markdown(tree_view_html(df, value_cols), unsafe_allow_html=True)
                st.markdown("### Summen")
                cols = st.columns(min(len(value_cols), 4))
                for i, c in enumerate(value_cols[:4]):
                    cols[i].metric(str(c), format_de_amount(df[c].sum(), "EUR"))

        with pivot_tab:
            level = st.slider("Aggregationsebene", min_value=1, max_value=7, value=5)
            pivot = build_pivot(df, level, value_cols)

            if pivot.empty:
                st.error("Keine auswertbare Ausweisstruktur gefunden.")
            else:
                st.dataframe(display_dataframe(pivot, value_cols), use_container_width=True, hide_index=True)


# ------------------------------------------------------------
# Phase 6
# ------------------------------------------------------------
elif phase == "7 Interpretation":
    st.subheader("KI-gestützte Interpretation")

    if st.session_state.mapped is None:
        st.warning("Bitte zuerst in Phase 4 das Mapping starten.")
    else:
        df = st.session_state.mapped.copy()
        value_cols = st.session_state.meta.get("value_cols", detect_value_cols(df))
        current_col, prior_col = analysis_columns(value_cols)

        if current_col is None:
            st.error("Keine Wertspalten gefunden.")
        else:
            control_1, control_2 = st.columns(2)
            with control_1:
                threshold = st.number_input(
                    "Wesentlichkeitsschwelle EUR",
                    min_value=0.0,
                    value=10000.0,
                    step=1000.0,
                )
            with control_2:
                top_n = st.slider("Anzahl Treiber", min_value=5, max_value=30, value=10)

            st.caption(f"Vergleich: {current_col} gegen {prior_col or 'kein Vorjahr'}")

            default_rules = [
                "Keine Tatsachen erfinden",
                "Hypothesen nur als prüfbedürftig formulieren",
                "HGB-konforme Sprache",
                "Wesentliche Auffälligkeiten priorisieren",
                "Keine Dopplungen",
                "Bei Unsicherheit Rückfragen formulieren",
            ]
            default_structure = ["Kernaussagen", "Wesentliche Treiber", "Risiken / Prüfhinweise", "Rückfragen"]

            level_summary = grouped_variance(df, ["Ausweis_1", "Ausweis_2", "Ausweis_3"], current_col, prior_col).head(top_n)
            account_work = add_variance_columns(df, current_col, prior_col)
            top_accounts = account_work[account_work["Veränderung_abs"] >= threshold].sort_values("Veränderung_abs", ascending=False).head(top_n)

            summary_tab, account_tab, prompt_tab, ai_tab = st.tabs(["Auffälligkeiten", "Kontentreiber", "KI-Arbeitsgrundlage", "OpenAI-Entwurf"])

            with summary_tab:
                st.markdown("### Auffällige Abschlusspositionen")
                if level_summary.empty:
                    st.info("Keine auswertbaren Abschlusspositionen gefunden.")
                else:
                    st.dataframe(variance_display(level_summary), use_container_width=True, hide_index=True)

            with account_tab:
                st.markdown("### Größte Kontentreiber")
                display_cols = ["KontoNr", "Kontobezeichnung", "Ausweis_1", "Ausweis_2", "Ausweis_3", "Laufendes Jahr", "Vorjahr", "Veränderung", "Veränderung_%"]
                display_cols = [c for c in display_cols if c in top_accounts.columns]
                if top_accounts.empty:
                    st.info("Keine Konten oberhalb der Schwelle.")
                else:
                    st.dataframe(variance_display(top_accounts[display_cols + ["Veränderung_abs"]]), use_container_width=True, hide_index=True)

            with prompt_tab:
                prompt_config = {
                    "role": "Wirtschaftsprüfer",
                    "purpose": "Management Summary",
                    "audience": "Geschäftsführung",
                    "tone": "prüferfreundlich, vorsichtig, sachlich, konservativ",
                    "length": "mittel",
                    "form": "gegliederte Abschnitte",
                    "rules": default_rules,
                    "output_structure": default_structure,
                }
                markdown = interpretation_markdown(
                    df,
                    value_cols,
                    st.session_state.mandant,
                    int(st.session_state.abschlussjahr),
                    threshold,
                    top_n,
                    prompt_config,
                    st.session_state.active_mandant,
                    st.session_state.reporting_profile,
                    st.session_state.onboarding_answers,
                )
                st.text_area("Prompt für KI-Interpretation", markdown, height=420)
                st.download_button(
                    "KI-Arbeitsgrundlage herunterladen",
                    data=markdown.encode("utf-8"),
                    file_name=f"Lumina_KI_Arbeitsgrundlage_{st.session_state.abschlussjahr}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            with ai_tab:
                st.markdown("### Steuerung")
                model_label = st.selectbox(
                    "OpenAI-Modell",
                    list(OPENAI_MODELS.keys()),
                    index=0,
                )
                model = OPENAI_MODELS[model_label]

                c_role, c_purpose = st.columns(2)
                with c_role:
                    role = st.selectbox("Rolle", AI_ROLES, index=0)
                with c_purpose:
                    purpose = st.selectbox("Zweck", AI_PURPOSES, index=2)

                c_audience, c_tone = st.columns(2)
                with c_audience:
                    audience = st.selectbox("Adressat", AI_AUDIENCES, index=0)
                with c_tone:
                    tone = st.selectbox("Ton / Stil", AI_TONES, index=0)

                c_temp, c_length, c_form = st.columns(3)
                with c_temp:
                    temperature = st.slider(
                        "Temperatur / Kreativität",
                        min_value=0.0,
                        max_value=1.0,
                        value=0.1,
                        step=0.1,
                    )
                with c_length:
                    text_length = st.selectbox("Textumfang", AI_LENGTHS, index=1)
                with c_form:
                    text_form = st.selectbox("Form", AI_FORMS, index=0)
                if not model.startswith(("gpt-4", "gpt-4o")):
                    st.caption("Hinweis: Die Kreativität wird bei GPT-5-Modellen vor allem über Rolle, Ton und Regeln gesteuert.")

                rules_text = st.text_area("Wichtige Regeln", "\n".join(default_rules), height=150)
                structure_text = st.text_area("Ausgabe-Struktur", "\n".join(default_structure), height=120)

                ai_config = {
                    "role": role,
                    "purpose": purpose,
                    "audience": audience,
                    "tone": tone,
                    "length": text_length,
                    "form": text_form,
                    "rules": [line.strip("- ").strip() for line in rules_text.splitlines() if line.strip()],
                    "output_structure": [line.strip("0123456789. -").strip() for line in structure_text.splitlines() if line.strip()],
                }

                markdown = interpretation_markdown(
                    df,
                    value_cols,
                    st.session_state.mandant,
                    int(st.session_state.abschlussjahr),
                    threshold,
                    top_n,
                    ai_config,
                    st.session_state.active_mandant,
                    st.session_state.reporting_profile,
                    st.session_state.onboarding_answers,
                )
                status = openai_status()
                if status == "verbunden":
                    st.success("OpenAI ist verbunden.")
                else:
                    st.warning(status)
                    st.caption("Lege den API-Key in Streamlit unter Secrets als OPENAI_API_KEY ab.")

                if st.button("Interpretation mit OpenAI erzeugen", type="primary", use_container_width=True):
                    progress = st.progress(0)
                    status_text = st.empty()

                    status_text.write("Schritt 1/4: Zahlenbasis und Auffälligkeiten werden vorbereitet.")
                    progress.progress(15)

                    status_text.write("Schritt 2/4: KI-Arbeitsgrundlage wird zusammengestellt.")
                    progress.progress(35)

                    status_text.write(f"Schritt 3/4: OpenAI wird mit Modell {model} angefragt. Das kann je nach Umfang etwas dauern.")
                    progress.progress(65)

                    with st.spinner("OpenAI erstellt den Interpretationsentwurf..."):
                        result, err = generate_openai_interpretation(markdown, model, temperature, ai_config)
                    if err:
                        progress.progress(100)
                        status_text.write("Die Anfrage wurde beendet, aber OpenAI hat einen Fehler zurückgegeben.")
                        st.error(err)
                    else:
                        st.session_state.ai_interpretation = result or ""
                        progress.progress(100)
                        status_text.write("Schritt 4/4: Entwurf wurde übernommen und steht unten bereit.")
                        st.success("Interpretationsentwurf erstellt.")

                if st.session_state.ai_interpretation:
                    st.markdown("### OpenAI-Entwurf")
                    st.markdown(st.session_state.ai_interpretation)
                    with st.expander("Textversion anzeigen"):
                        st.text_area("OpenAI-Entwurf als Text", st.session_state.ai_interpretation, height=420)
                    docx = ai_output_to_docx(
                        st.session_state.ai_interpretation,
                        f"LUMINA Interpretation {st.session_state.mandant} {st.session_state.abschlussjahr}",
                    )
                    if docx:
                        st.download_button(
                            "OpenAI-Entwurf als Word herunterladen",
                            data=docx,
                            file_name=f"Lumina_OpenAI_Interpretation_{st.session_state.abschlussjahr}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    else:
                        st.warning("Word-Export ist erst verfügbar, wenn python-docx installiert ist.")


# ------------------------------------------------------------
# Phase 8
# ------------------------------------------------------------
elif phase == "8 Export":
    st.subheader("Export & Prüfungspaket")

    if st.session_state.mapped is None:
        st.warning("Bitte zuerst in Phase 4 das Mapping starten.")
    else:
        df = st.session_state.mapped.copy()
        value_cols = st.session_state.meta.get("value_cols", detect_value_cols(df))
        pivot = build_pivot(df, 5, value_cols)
        klarung = df[clarification_mask(df)].copy()

        st.markdown(
            f"""
            <div class="lumina-box">
            <b>Exportinhalt</b><br>
            Mandant: {st.session_state.mandant}<br>
            Abschlussjahr: {st.session_state.abschlussjahr}<br>
            Standard: {st.session_state.standard}<br>
            Klärungsposten: {len(klarung)}
            </div>
            """,
            unsafe_allow_html=True,
        )

        xlsx = excel_export(
            st.session_state.susa_raw if st.session_state.susa_raw is not None else st.session_state.susa_norm,
            df,
            pivot,
            klarung,
        )
        safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", st.session_state.mandant).strip("_")
        filename = f"Lumina_Pruefungspaket_{safe_name}_{st.session_state.abschlussjahr}.xlsx"
        if st.session_state.export_status != "bereit":
            st.session_state.export_status = "bereit"
            audit_log("export_created", f"Excel-Prüfungspaket vorbereitet: {filename}")

        st.download_button(
            "📥 Excel-Prüfungspaket herunterladen",
            data=xlsx,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            type="primary",
            use_container_width=True,
        )

        st.info("PDF würde ich erst später ergänzen. Für Wirtschaftsprüfer ist zuerst ein sauberer Excel-Export wertvoller.")
